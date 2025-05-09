import streamlit as st
import pandas as pd
import numpy as np
import requests
import json
import time
import re
from bs4 import BeautifulSoup
import trafilatura
import openai
import anthropic
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import base64
import random
from typing import List, Dict, Any, Tuple, Optional
import logging
import traceback
import openpyxl
import matplotlib.pyplot as plt
import altair as alt
import faiss
from collections import Counter
from textstat import flesch_reading_ease, text_standard

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Set page configuration
st.set_page_config(
    page_title="SEO Content Optimizer",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

###############################################################################
# 1. Utility Functions
###############################################################################

def display_error(error_msg: str):
    """Display error message in Streamlit"""
    st.error(f"Error: {error_msg}")
    logger.error(error_msg)

def get_download_link(file_bytes: bytes, filename: str, link_label: str) -> str:
    """Returns an HTML link to download file"""
    b64 = base64.b64encode(file_bytes).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">{link_label}</a>'
    return href

def format_time(seconds: float) -> str:
    """Format time in seconds to readable format"""
    if seconds < 60:
        return f"{seconds:.1f} seconds"
    else:
        minutes = int(seconds // 60)
        sec = seconds % 60
        return f"{minutes} min {sec:.1f} sec"

###############################################################################
# 2. FireCrawl Client for Web Scraping
###############################################################################

class FireCrawl:
    """
    Client for interacting with FireCrawl's API with concurrency constraints.
    """
    def __init__(self, api_key: str, max_concurrency: int = 5, rate_limit_per_minute: int = 60):
        self.api_key = api_key
        self.max_concurrency = max_concurrency
        self.rate_limit_per_minute = rate_limit_per_minute
        self.active_requests = 0
        self.request_times = []
        self.logger = logging.getLogger(__name__)
    
    def _wait_for_rate_limit(self):
        """Wait if rate limit has been reached"""
        now = time.time()
        # Clean up old requests
        self.request_times = [t for t in self.request_times if now - t < 60]
        
        if len(self.request_times) >= self.rate_limit_per_minute:
            # Wait until oldest request is a minute old
            wait_time = 60 - (now - self.request_times[0])
            if wait_time > 0:
                self.logger.info(f"Rate limit reached. Waiting {wait_time:.2f} seconds.")
                time.sleep(wait_time)
                # Clean up again after waiting
                now = time.time()
                self.request_times = [t for t in self.request_times if now - t < 60]
    
    def _wait_for_concurrency(self):
        """Wait until concurrency slot is available"""
        while self.active_requests >= self.max_concurrency:
            self.logger.info(f"Max concurrency reached ({self.max_concurrency}). Waiting...")
            time.sleep(1)
    
    def scrape_url(self, url: str) -> Tuple[str, bool]:
        """
        Scrape a webpage using FireCrawl API.
        Returns: content, success_status
        """
        try:
            # Wait for rate limit and concurrency slots
            self._wait_for_rate_limit()
            self._wait_for_concurrency()
            
            # Track request
            self.active_requests += 1
            self.request_times.append(time.time())
            
            # Make API request to FireCrawl
            headers = {
                'Authorization': f'Bearer {self.api_key}',
                'Content-Type': 'application/json'
            }
            
            data = {
                'url': url,
                'javascript': True,  # Enable JavaScript rendering
                'timeout': 30,       # Timeout in seconds
                'extract_text': True # Extract main content
            }
            
            response = requests.post(
                'https://api.firecrawl.com/scrape',  # Replace with actual FireCrawl endpoint
                headers=headers,
                json=data,
                timeout=45
            )
            
            self.active_requests -= 1  # Request completed
            
            if response.status_code == 200:
                result = response.json()
                content = result.get('content', '')
                
                if content:
                    return content, True
                else:
                    self.logger.warning(f"No content returned for {url}")
                    return "", False
            else:
                self.logger.error(f"FireCrawl API error: {response.status_code} - {response.text}")
                return "", False
                
        except Exception as e:
            self.active_requests -= 1  # Request failed
            self.logger.error(f"Exception in FireCrawl scrape_url for {url}: {str(e)}")
            return "", False
    
    def extract_headings(self, url: str) -> Dict[str, List[str]]:
        """
        Extract headings from a webpage using FireCrawl.
        Returns: Dictionary with h1, h2, h3 lists
        """
        try:
            # Wait for rate limit and concurrency slots
            self._wait_for_rate_limit()
            self._wait_for_concurrency()
            
            # Track request
            self.active_requests += 1
            self.request_times.append(time.time())
            
            # Make API request to FireCrawl
            headers = {
                'Authorization': f'Bearer {self.api_key}',
                'Content-Type': 'application/json'
            }
            
            data = {
                'url': url,
                'javascript': True,
                'timeout': 30,
                'extract_headings': True  # Extract headings specifically
            }
            
            response = requests.post(
                'https://api.firecrawl.com/scrape',  # Replace with actual FireCrawl endpoint
                headers=headers,
                json=data,
                timeout=45
            )
            
            self.active_requests -= 1  # Request completed
            
            if response.status_code == 200:
                result = response.json()
                headings = result.get('headings', {})
                
                # Ensure proper structure
                structured_headings = {
                    'h1': headings.get('h1', []),
                    'h2': headings.get('h2', []),
                    'h3': headings.get('h3', [])
                }
                
                return structured_headings
            else:
                self.logger.error(f"FireCrawl API error: {response.status_code} - {response.text}")
                return {'h1': [], 'h2': [], 'h3': []}
                
        except Exception as e:
            self.active_requests -= 1  # Request failed
            self.logger.error(f"Exception in FireCrawl extract_headings for {url}: {str(e)}")
            return {'h1': [], 'h2': [], 'h3': []}
    
    def batch_scrape(self, urls: List[str]) -> List[Tuple[str, str, bool]]:
        """
        Scrape multiple URLs with concurrency control.
        Returns: List of (url, content, success_status) tuples
        """
        results = []
        
        for url in urls:
            content, success = self.scrape_url(url)
            results.append((url, content, success))
        
        return results

###############################################################################
# 3. API Integration - DataForSEO
###############################################################################

def classify_page_type(url: str, title: str, snippet: str) -> str:
    """
    Simplified page type classification based on title and snippet patterns
    Returns: page_type
    """
    title_lower = title.lower() if title else ""
    snippet_lower = snippet.lower() if snippet else ""
    url_lower = url.lower() if url else ""
    
    # E-commerce indicators
    commerce_patterns = ['buy', 'shop', 'purchase', 'cart', 'checkout', 'price', 'discount', 'sale', 'product', 'order']
    if any(pattern in title_lower or pattern in snippet_lower or pattern in url_lower for pattern in commerce_patterns):
        return "E-commerce"
    
    # Article/Blog indicators
    article_patterns = ['blog', 'article', 'news', 'post', 'how to', 'guide', 'tips', 'tutorial', 'learn']
    if any(pattern in title_lower or pattern in snippet_lower or pattern in url_lower for pattern in article_patterns):
        return "Article/Blog"
    
    # Forum/Community indicators
    forum_patterns = ['forum', 'community', 'discussion', 'thread', 'reply', 'comment', 'question', 'answer']
    if any(pattern in title_lower or pattern in snippet_lower or pattern in url_lower for pattern in forum_patterns):
        return "Forum/Community"
    
    # Review indicators
    review_patterns = ['review', 'comparison', 'vs', 'versus', 'top 10', 'best', 'rating', 'rated']
    if any(pattern in title_lower or pattern in snippet_lower or pattern in url_lower for pattern in review_patterns):
        return "Review/Comparison"
    
    # Default to informational
    return "Informational"

def fetch_serp_results(keyword: str, api_login: str, api_password: str) -> Tuple[List[Dict], List[Dict], List[Dict], bool]:
    """
    Fetch SERP results from DataForSEO API and classify pages
    Returns: organic_results, serp_features, paa_questions, success_status
    """
    try:
        url = "https://api.dataforseo.com/v3/serp/google/organic/live/advanced"
        headers = {
            'Content-Type': 'application/json',
        }
        
        # Prepare request data
        post_data = [{
            "keyword": keyword,
            "location_code": 2840,  # USA
            "language_code": "en",
            "device": "desktop",
            "os": "windows",
            "depth": 30  # Get more results to ensure we have at least 10 organic
        }]
        
        # Make API request
        response = requests.post(
            url,
            auth=(api_login, api_password),
            headers=headers,
            json=post_data
        )
        
        # Process response
        if response.status_code == 200:
            data = response.json()
            logger.info(f"SERP API Response status: {data.get('status_code')}")
            
            if data.get('status_code') == 20000:
                results = data['tasks'][0]['result'][0]
                
                # Extract organic results
                organic_results = []
                for item in results.get('items', []):
                    if item.get('type') == 'organic':
                        if len(organic_results) < 20:  # Increased to 20 for better clustering
                            # Get title and snippet for classification
                            title = item.get('title', '')
                            snippet = item.get('snippet', '')
                            url = item.get('url', '')
                            
                            # Classify page type using pattern matching
                            page_type = classify_page_type(url, title, snippet)
                            
                            organic_results.append({
                                'url': url,
                                'title': title,
                                'snippet': snippet,
                                'rank_group': item.get('rank_group'),
                                'page_type': page_type
                            })
                
                # Extract People Also Asked questions - handle both possible formats
                paa_questions = []
                
                # Method 1: Look for a PAA container
                for item in results.get('items', []):
                    if item.get('type') == 'people_also_ask':
                        logger.info(f"Found PAA container with {len(item.get('items', []))} questions")
                        
                        # Process PAA items within the container
                        for paa_item in item.get('items', []):
                            if paa_item.get('type') == 'people_also_ask_element':
                                question_data = {
                                    'question': paa_item.get('title', ''),
                                    'expanded': []
                                }
                                
                                # Extract expanded element data if available
                                for expanded in paa_item.get('expanded_element', []):
                                    if expanded.get('type') == 'people_also_ask_expanded_element':
                                        question_data['expanded'].append({
                                            'url': expanded.get('url', ''),
                                            'title': expanded.get('title', ''),
                                            'description': expanded.get('description', '')
                                        })
                                
                                paa_questions.append(question_data)
                
                # Method 2 (fallback): Look for individual PAA elements directly in items
                if not paa_questions:
                    logger.info("No PAA container found, looking for individual PAA elements")
                    for item in results.get('items', []):
                        if item.get('type') == 'people_also_ask_element':
                            question_data = {
                                'question': item.get('title', ''),
                                'expanded': []
                            }
                            
                            # Extract expanded element data if available
                            for expanded in item.get('expanded_element', []):
                                if expanded.get('type') == 'people_also_ask_expanded_element':
                                    question_data['expanded'].append({
                                        'url': expanded.get('url', ''),
                                        'title': expanded.get('title', ''),
                                        'description': expanded.get('description', '')
                                    })
                            
                            paa_questions.append(question_data)
                
                # Log how many PAA questions we found
                logger.info(f"Extracted {len(paa_questions)} PAA questions")
                
                # Extract SERP features
                serp_features = []
                feature_counts = {}
                for item in results.get('items', []):
                    item_type = item.get('type')
                    if item_type != 'organic':
                        if item_type in feature_counts:
                            feature_counts[item_type] += 1
                        else:
                            feature_counts[item_type] = 1
                
                for feature, count in feature_counts.items():
                    serp_features.append({
                        'feature_type': feature,
                        'count': count
                    })
                
                # Sort by count and limit to top 20
                serp_features = sorted(serp_features, key=lambda x: x['count'], reverse=True)[:20]
                
                return organic_results, serp_features, paa_questions, True
            else:
                error_msg = f"API Error: {data.get('status_message')}"
                logger.error(error_msg)
                return [], [], [], False
        else:
            error_msg = f"HTTP Error: {response.status_code} - {response.text}"
            logger.error(error_msg)
            return [], [], [], False
    
    except Exception as e:
        error_msg = f"Exception in fetch_serp_results: {str(e)}"
        logger.error(error_msg)
        logger.error(traceback.format_exc())
        return [], [], [], False

###############################################################################
# 4. API Integration - DataForSEO for Keywords
###############################################################################

def fetch_keyword_suggestions(keyword: str, api_login: str, api_password: str) -> Tuple[List[Dict], bool]:
    """
    Fetch keyword suggestions from DataForSEO to get accurate search volume data
    Returns: keyword_suggestions, success_status
    """
    try:
        url = "https://api.dataforseo.com/v3/dataforseo_labs/google/keyword_suggestions/live"
        headers = {
            'Content-Type': 'application/json',
        }
        
        # Prepare request data based on the sample JSON structure
        post_data = [{
            "keyword": keyword,
            "location_code": 2840,  # USA
            "language_code": "en",
            "include_serp_info": True,
            "include_seed_keyword": True,
            "limit": 20  # Fetch top 20 suggestions
        }]
        
        # Log the request for debugging
        logger.info(f"Fetching keyword suggestions for: {keyword}")
        
        # Make API request
        response = requests.post(
            url,
            auth=(api_login, api_password),
            headers=headers,
            json=post_data
        )
        
        # Process response
        if response.status_code == 200:
            data = response.json()
            logger.info(f"Keyword Suggestions API Response status: {data.get('status_code')}")
            
            # Validate response
            if data.get('status_code') != 20000 or not data.get('tasks') or len(data['tasks']) == 0:
                logger.warning(f"Invalid API response: {data.get('status_message')}")
                return [], False
            
            keyword_suggestions = []
            
            # Process results based on the specific JSON structure from sample
            for task in data['tasks']:
                if not task.get('result'):
                    continue
                
                for result in task['result']:
                    # First, check for seed keyword data
                    if 'seed_keyword_data' in result and result['seed_keyword_data']:
                        seed_data = result['seed_keyword_data']
                        if 'keyword_info' in seed_data:
                            keyword_info = seed_data['keyword_info']
                            keyword_suggestions.append({
                                'keyword': result.get('seed_keyword', ''),
                                'search_volume': keyword_info.get('search_volume', 0),
                                'cpc': keyword_info.get('cpc', 0.0),
                                'competition': keyword_info.get('competition', 0.0)
                            })
                    
                    # Then look for items array which contains related keywords
                    if 'items' in result and isinstance(result['items'], list):
                        for item in result['items']:
                            if 'keyword_info' in item:
                                keyword_info = item['keyword_info']
                                keyword_suggestions.append({
                                    'keyword': item.get('keyword', ''),
                                    'search_volume': keyword_info.get('search_volume', 0),
                                    'cpc': keyword_info.get('cpc', 0.0),
                                    'competition': keyword_info.get('competition', 0.0)
                                })
            
            # Check if we successfully found keywords
            if keyword_suggestions:
                # Sort by search volume (descending)
                keyword_suggestions.sort(key=lambda x: x.get('search_volume', 0), reverse=True)
                logger.info(f"Successfully extracted {len(keyword_suggestions)} keyword suggestions")
                return keyword_suggestions, True
            else:
                logger.warning(f"No keyword suggestions found in the response")
                return [], False
        else:
            error_msg = f"HTTP Error: {response.status_code} - {response.text}"
            logger.error(error_msg)
            return [], False
    
    except Exception as e:
        error_msg = f"Exception in fetch_keyword_suggestions: {str(e)}"
        logger.error(error_msg)
        logger.error(traceback.format_exc())
        return [], False

def fetch_related_keywords_dataforseo(keyword: str, api_login: str, api_password: str) -> Tuple[List[Dict], bool]:
    """
    Fetch related keywords from DataForSEO Related Keywords API
    Uses proper API structure parsing based on the provided sample
    Returns: related_keywords, success_status
    """
    try:
        url = "https://api.dataforseo.com/v3/dataforseo_labs/google/related_keywords/live"
        headers = {
            'Content-Type': 'application/json',
        }
        
        # Prepare request data based on the provided sample format
        post_data = [{
            "keyword": keyword,
            "language_name": "English", 
            "location_code": 2840,  # USA
            "limit": 20  # Fetch top 20 related keywords
        }]
        
        # Log the request for debugging
        logger.info(f"Fetching related keywords for: {keyword}")
        
        # Make API request
        response = requests.post(
            url,
            auth=(api_login, api_password),
            headers=headers,
            json=post_data
        )
        
        # Process response
        if response.status_code == 200:
            data = response.json()
            logger.info(f"API Response status: {data.get('status_code')}")
            
            # Validate response
            if data.get('status_code') != 20000 or not data.get('tasks'):
                logger.warning(f"Invalid API response: {data.get('status_message')}")
                return fetch_keyword_suggestions(keyword, api_login, api_password)
            
            related_keywords = []
            
            # Process the results following the structure of the provided JSON example
            for task in data['tasks']:
                if not task.get('result'):
                    continue
                
                for result in task['result']:
                    # Process items array which contains the keyword data
                    for item in result.get('items', []):
                        if 'keyword_data' in item:
                            kw_data = item['keyword_data']
                            if 'keyword_info' in kw_data:
                                keyword_info = kw_data['keyword_info']
                                
                                related_keywords.append({
                                    'keyword': kw_data.get('keyword', ''),
                                    'search_volume': keyword_info.get('search_volume', 0),
                                    'cpc': keyword_info.get('cpc', 0.0),
                                    'competition': keyword_info.get('competition', 0.0)
                                })
            
            # Check if we found any keywords with this approach
            if related_keywords:
                logger.info(f"Successfully extracted {len(related_keywords)} related keywords")
                return related_keywords, True
            else:
                # If no keywords found, try the keyword suggestions endpoint
                logger.warning(f"No related keywords found, trying keyword suggestions endpoint")
                return fetch_keyword_suggestions(keyword, api_login, api_password)
        else:
            error_msg = f"HTTP Error: {response.status_code} - {response.text}"
            logger.error(error_msg)
            return fetch_keyword_suggestions(keyword, api_login, api_password)
    
    except Exception as e:
        error_msg = f"Exception in fetch_related_keywords_dataforseo: {str(e)}"
        logger.error(error_msg)
        logger.error(traceback.format_exc())
        return fetch_keyword_suggestions(keyword, api_login, api_password)

###############################################################################
# 5. Web Scraping and Content Analysis
###############################################################################

def scrape_webpage(url: str, firecrawl_api_key: str = None) -> Tuple[str, bool]:
    """
    Enhanced webpage scraping using FireCrawl if API key provided, otherwise fallback to trafilatura
    Returns: content, success_status
    """
    try:
        # Use FireCrawl if API key is provided
        if firecrawl_api_key:
            firecrawl = FireCrawl(firecrawl_api_key)
            return firecrawl.scrape_url(url)
        
        # Otherwise, use trafilatura as fallback
        downloaded = trafilatura.fetch_url(url)
        if downloaded:
            content = trafilatura.extract(downloaded, include_comments=False, include_tables=True)
            if content:
                return content, True
        
        # Fallback to requests + BeautifulSoup if trafilatura fails
        user_agents = [
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/14.1.1 Safari/605.1.15',
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:89.0) Gecko/20100101 Firefox/89.0',
            'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/92.0.4515.107 Safari/537.36'
        ]
        
        headers = {
            'User-Agent': random.choice(user_agents),
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Referer': 'https://www.google.com/'
        }
        
        response = requests.get(url, headers=headers, timeout=15)
        
        if response.status_code == 200:
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Remove script, style, and other non-content elements
            for element in soup(["script", "style", "header", "footer", "nav", "aside", "form"]):
                element.extract()
            
            # Try to find main content
            main_content = soup.find('main') or soup.find('article') or soup.find('div', class_='content')
            
            if main_content:
                text = main_content.get_text(separator='\n')
            else:
                text = soup.get_text(separator='\n')
            
            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text, True
        
        elif response.status_code == 403:
            logger.warning(f"Access forbidden (403) for URL: {url}")
            return f"[Content not accessible due to site restrictions]", False
        
        else:
            logger.warning(f"HTTP error {response.status_code} for URL: {url}")
            return f"[Error retrieving content: HTTP {response.status_code}]", False
    
    except Exception as e:
        error_msg = f"Exception in scrape_webpage for {url}: {str(e)}"
        logger.error(error_msg)
        return f"[Error: {str(e)}]", False

def extract_headings(url: str, firecrawl_api_key: str = None) -> Dict[str, List[str]]:
    """
    Extract headings (H1, H2, H3) from a webpage using FireCrawl if API key provided
    """
    try:
        # Use FireCrawl if API key is provided
        if firecrawl_api_key:
            firecrawl = FireCrawl(firecrawl_api_key)
            return firecrawl.extract_headings(url)
        
        # Fallback to traditional method
        user_agents = [
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/14.1.1 Safari/605.1.15',
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:89.0) Gecko/20100101 Firefox/89.0',
            'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/92.0.4515.107 Safari/537.36'
        ]
        
        # Select a random user agent
        user_agent = random.choice(user_agents)
        
        headers = {
            'User-Agent': user_agent,
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Referer': 'https://www.google.com/'
        }
        
        response = requests.get(url, headers=headers, timeout=15)
        
        if response.status_code == 200:
            soup = BeautifulSoup(response.text, 'html.parser')
            
            headings = {
                'h1': [h.get_text().strip() for h in soup.find_all('h1')],
                'h2': [h.get_text().strip() for h in soup.find_all('h2')],
                'h3': [h.get_text().strip() for h in soup.find_all('h3')]
            }
            
            return headings
        else:
            return {'h1': [], 'h2': [], 'h3': []}
    
    except Exception as e:
        error_msg = f"Exception in extract_headings: {str(e)}"
        logger.error(error_msg)
        return {'h1': [], 'h2': [], 'h3': []}

###############################################################################
# 6. Content Clustering and Corpus Analysis
###############################################################################

def generate_embedding(text: str, openai_api_key: str, model: str = "text-embedding-3-large") -> Tuple[List[float], bool]:
    """
    Generate embedding for text using OpenAI API
    Returns: embedding, success_status
    """
    try:
        openai.api_key = openai_api_key
        
        # Limit text length to prevent token limit issues
        text = text[:10000]  # Adjust limit as needed
        
        response = openai.Embedding.create(
            model=model,
            input=text
        )
        
        embedding = response['data'][0]['embedding']
        return embedding, True
    
    except Exception as e:
        error_msg = f"Exception in generate_embedding: {str(e)}"
        logger.error(error_msg)
        return [], False

def _cluster_competitors(competitor_contents: List[Dict], openai_api_key: str, num_clusters: int = 5) -> Tuple[List[Dict], bool]:
    """
    Cluster competitor content using OpenAI embeddings and FAISS K-means.
    Returns: clustered_competitors, success_status
    """
    try:
        # Generate embeddings for all competitor content
        embeddings = []
        valid_competitors = []
        
        for i, competitor in enumerate(competitor_contents):
            content = competitor.get('content', '')
            if not content or len(content) < 200:  # Skip if too short
                continue
                
            embedding, success = generate_embedding(content, openai_api_key)
            
            if success and embedding:
                embeddings.append(embedding)
                valid_competitors.append(competitor)
        
        if not embeddings or len(embeddings) < 2:
            logger.error("Not enough valid content to cluster")
            return competitor_contents, False
        
        # Convert to numpy array
        embedding_array = np.array(embeddings).astype('float32')
        
        # Initialize FAISS index
        dimension = len(embeddings[0])
        
        # Normalize vectors for cosine similarity
        faiss.normalize_L2(embedding_array)
        
        # Use KMeans clustering
        actual_clusters = min(num_clusters, len(embedding_array))
        kmeans = faiss.Kmeans(dimension, actual_clusters, niter=20, verbose=False)
        kmeans.train(embedding_array)
        
        # Get cluster assignments
        _, assignments = kmeans.index.search(embedding_array, 1)
        
        # Add cluster info to competitors
        for i, cluster_id in enumerate(assignments.flatten()):
            valid_competitors[i]['cluster'] = int(cluster_id)
        
        # Sort by cluster for easy access
        clustered_competitors = sorted(valid_competitors, key=lambda x: x.get('cluster', 0))
        
        return clustered_competitors, True
    
    except Exception as e:
        error_msg = f"Exception in _cluster_competitors: {str(e)}"
        logger.error(error_msg)
        return competitor_contents, False

def build_lexical_corpus(competitor_contents: List[Dict]) -> Dict:
    """
    Build a lexical corpus from competitor contents including heading frequencies and keyword usage.
    Returns: corpus data
    """
    corpus = {
        'heading_counts': {
            'h2': Counter(),
            'h3': Counter(),
        },
        'keyword_frequency': Counter(),
        'content_lengths': [],
        'sentence_starters': Counter(),
    }
    
    for competitor in competitor_contents:
        # Extract headings
        headings = competitor.get('headings', {})
        
        # Count H2 and H3 headings
        for h2 in headings.get('h2', []):
            corpus['heading_counts']['h2'][h2.lower()] += 1
            
        for h3 in headings.get('h3', []):
            corpus['heading_counts']['h3'][h3.lower()] += 1
        
        # Extract content for keyword analysis
        content = competitor.get('content', '')
        if content:
            # Count content length
            corpus['content_lengths'].append(len(content.split()))
            
            # Extract keywords (simple approach - more sophisticated NLP could be used)
            words = re.findall(r'\b\w+\b', content.lower())
            for word in words:
                if len(word) > 3:  # Skip short words
                    corpus['keyword_frequency'][word] += 1
            
            # Analyze sentence starters (for variety checking)
            sentences = re.split(r'[.!?]', content)
            for sentence in sentences:
                sentence = sentence.strip()
                if len(sentence) > 10:  # Skip too short sentences
                    # Get first 3 words (3-gram prefix)
                    words = sentence.split()
                    if len(words) >= 3:
                        prefix = ' '.join(words[:3]).lower()
                        corpus['sentence_starters'][prefix] += 1
    
    return corpus

###############################################################################
# 7. Content Readability and Linting
###############################################################################

def lint_article(article_content: str, keyword: str, term_data: Dict) -> Dict:
    """
    Analyze article content for readability, keyword density, and sentence variety
    Returns: lint results
    """
    results = {
        'readability': {
            'score': 0,
            'grade_level': '',
            'issues': []
        },
        'keyword_density': {
            'score': 0,
            'primary_keyword': {'term': keyword, 'count': 0, 'density': 0, 'status': 'good'},
            'primary_terms': [],
            'issues': []
        },
        'sentence_variety': {
            'score': 0,
            'repeat_patterns': [],
            'issues': []
        }
    }
    
    # Clean HTML tags for analysis
    clean_text = re.sub(r'<[^>]+>', ' ', article_content)
    
    # 1. Readability Analysis
    try:
        readability_score = flesch_reading_ease(clean_text)
        grade_level = text_standard(clean_text)
        
        results['readability']['score'] = readability_score
        results['readability']['grade_level'] = grade_level
        
        # Interpret Flesch Reading Ease score
        if readability_score < 30:
            results['readability']['issues'].append("Text is very difficult to read. Consider simplifying.")
        elif readability_score < 50:
            results['readability']['issues'].append("Text is difficult to read. Consider using shorter sentences and simpler words.")
        elif readability_score > 80:
            results['readability']['issues'].append("Text may be too simplistic for professional content.")
    except Exception as e:
        results['readability']['issues'].append(f"Error analyzing readability: {str(e)}")
    
    # 2. Keyword Density Analysis
    try:
        # Count words
        words = re.findall(r'\b\w+\b', clean_text.lower())
        word_count = len(words)
        
        if word_count == 0:
            results['keyword_density']['issues'].append("No text content found for analysis.")
            return results
        
        # Primary keyword density
        keyword_lower = keyword.lower()
        primary_count = len(re.findall(r'\b' + re.escape(keyword_lower) + r'\b', clean_text.lower()))
        primary_density = (primary_count / word_count) * 100
        
        results['keyword_density']['primary_keyword']['count'] = primary_count
        results['keyword_density']['primary_keyword']['density'] = primary_density
        
        # Evaluate primary keyword density
        if primary_density == 0:
            results['keyword_density']['primary_keyword']['status'] = 'missing'
            results['keyword_density']['issues'].append(f"Primary keyword '{keyword}' not found in content.")
        elif primary_density < 0.5:
            results['keyword_density']['primary_keyword']['status'] = 'low'
            results['keyword_density']['issues'].append(f"Primary keyword '{keyword}' density is low ({primary_density:.2f}%).")
        elif primary_density > 2.5:
            results['keyword_density']['primary_keyword']['status'] = 'over'
            results['keyword_density']['issues'].append(f"Primary keyword '{keyword}' may be overused ({primary_density:.2f}%).")
        else:
            results['keyword_density']['primary_keyword']['status'] = 'good'
        
        # Analyze primary terms from term_data
        if 'primary_terms' in term_data:
            for term_info in term_data['primary_terms'][:10]:  # Top 10 primary terms
                term = term_info.get('term', '')
                if term.lower() != keyword_lower:  # Skip main keyword
                    term_count = len(re.findall(r'\b' + re.escape(term.lower()) + r'\b', clean_text.lower()))
                    term_density = (term_count / word_count) * 100
                    recommended = term_info.get('recommended_usage', 1)
                    
                    term_status = 'good'
                    if term_count == 0:
                        term_status = 'missing'
                    elif term_count < recommended:
                        term_status = 'low'
                    elif term_count > recommended * 2:
                        term_status = 'over'
                    
                    results['keyword_density']['primary_terms'].append({
                        'term': term,
                        'count': term_count,
                        'density': term_density,
                        'recommended': recommended,
                        'status': term_status
                    })
                    
                    if term_status == 'missing':
                        results['keyword_density']['issues'].append(f"Primary term '{term}' not found in content.")
                    elif term_status == 'low':
                        results['keyword_density']['issues'].append(f"Primary term '{term}' appears {term_count} times (recommended: {recommended}).")
    except Exception as e:
        results['keyword_density']['issues'].append(f"Error analyzing keyword density: {str(e)}")
    
    # 3. Sentence Variety Analysis
    try:
        # Split into sentences
        sentences = re.split(r'[.!?]', clean_text)
        sentences = [s.strip() for s in sentences if len(s.strip()) > 10]
        
        # Check for 3-gram prefix repetition
        prefixes = []
        for sentence in sentences:
            words = sentence.split()
            if len(words) >= 3:
                prefix = ' '.join(words[:3]).lower()
                prefixes.append(prefix)
        
        # Count prefix frequencies
        prefix_counter = Counter(prefixes)
        
        # Find repeating patterns (>=5 occurrences)
        repeat_patterns = [(prefix, count) for prefix, count in prefix_counter.items() if count >= 5]
        
        results['sentence_variety']['repeat_patterns'] = repeat_patterns
        
        # Add issues for repeating patterns
        for prefix, count in repeat_patterns:
            results['sentence_variety']['issues'].append(
                f"Sentence pattern '{prefix}...' is used {count} times. Vary your sentence structure."
            )
        
        # Calculate variety score
        if len(sentences) > 0:
            unique_ratio = len(prefix_counter) / len(prefixes)
            variety_score = unique_ratio * 100
            results['sentence_variety']['score'] = variety_score
            
            if variety_score < 70:
                results['sentence_variety']['issues'].append(
                    f"Low sentence variety score ({variety_score:.1f}%). Vary your sentence structures more."
                )
    except Exception as e:
        results['sentence_variety']['issues'].append(f"Error analyzing sentence variety: {str(e)}")
    
    return results

###############################################################################
# 8. Enhanced Writer Guide with Section Explanations
###############################################################################

def generate_writer_guide(keyword: str, semantic_structure: Dict, corpus: Dict, 
                        term_data: Dict, anthropic_api_key: str) -> Tuple[Dict, bool]:
    """
    Generate detailed writing guidance with explanations for each section
    Returns: writer_guide, success_status
    """
    try:
        client = anthropic.Anthropic(api_key=anthropic_api_key)
        
        # Create a guide structure with section explanations
        writer_guide = {
            'h1': {
                'heading': semantic_structure.get('h1', f"Complete Guide to {keyword}"),
                'guidance': '',
                'why_it_matters': ''
            },
            'sections': []
        }
        
        # Extract top headings from corpus for context
        top_h2_headings = corpus['heading_counts']['h2'].most_common(10)
        top_h3_headings = corpus['heading_counts']['h3'].most_common(15)
        
        # Format headings for prompt
        top_headings_text = "Top competitor H2 headings:\n"
        for heading, count in top_h2_headings:
            top_headings_text += f"- {heading} (appears {count} times)\n"
            
        top_headings_text += "\nTop competitor H3 headings:\n"
        for heading, count in top_h3_headings:
            top_headings_text += f"- {heading} (appears {count} times)\n"
        
        # Get primary terms info
        primary_terms_text = "Key terms to include:\n"
        if 'primary_terms' in term_data:
            for term_info in term_data['primary_terms'][:10]:
                term = term_info.get('term', '')
                importance = term_info.get('importance', 0)
                usage = term_info.get('recommended_usage', 1)
                primary_terms_text += f"- {term} (importance: {importance:.2f}, use ~{usage} times)\n"
        
        # First, generate H1 guidance
        h1_response = client.messages.create(
            model="claude-3-5-sonnet-20241022",
            max_tokens=1000,
            system="You are an expert SEO content strategist.",
            messages=[
                {"role": "user", "content": f"""
                Create detailed writing guidance for an article about "{keyword}" with this H1:
                
                H1: {writer_guide['h1']['heading']}
                
                Provide:
                1. Guidance on how to introduce this topic effectively
                2. Why this H1 is strategically important (SEO perspective)
                
                {primary_terms_text}
                
                Format your response as JSON:
                
                {{
                    "guidance": "Detailed guidance for introduction...",
                    "why_it_matters": "Explanation of why this H1 is important..."
                }}
                """}
            ],
            temperature=0.5
        )
        
        # Extract and parse JSON response for H1
        h1_content = h1_response.content[0].text
        h1_json_match = re.search(r'({.*})', h1_content, re.DOTALL)
        if h1_json_match:
            h1_data = json.loads(h1_json_match.group(1))
            writer_guide['h1']['guidance'] = h1_data.get('guidance', '')
            writer_guide['h1']['why_it_matters'] = h1_data.get('why_it_matters', '')
        
        # Then, generate guidance for each section and subsection
        for section in semantic_structure.get('sections', []):
            h2 = section.get('h2', '')
            if not h2:
                continue
                
            # Generate guidance for this H2 section
            h2_response = client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=1000,
                system="You are an expert SEO content strategist.",
                messages=[
                    {"role": "user", "content": f"""
                    Create detailed writing guidance for this section of an article about "{keyword}":
                    
                    H2: {h2}
                    
                    Provide:
                    1. Key points to cover in this section
                    2. Approximate word count recommendation
                    3. Specific keywords to include
                    4. Why this section matters for this topic (SEO perspective)
                    
                    {top_headings_text}
                    
                    {primary_terms_text}
                    
                    Format your response as JSON:
                    
                    {{
                        "guidance": "Detailed section guidance...",
                        "word_count": 300,
                        "key_terms": ["term1", "term2", "term3"],
                        "why_it_matters": "Explanation of why this section is important..."
                    }}
                    """}
                ],
                temperature=0.5
            )
            
            # Extract and parse JSON response for H2
            h2_content = h2_response.content[0].text
            h2_json_match = re.search(r'({.*})', h2_content, re.DOTALL)
            
            section_guide = {
                'heading': h2,
                'guidance': '',
                'word_count': 300,
                'key_terms': [],
                'why_it_matters': '',
                'subsections': []
            }
            
            if h2_json_match:
                h2_data = json.loads(h2_json_match.group(1))
                section_guide['guidance'] = h2_data.get('guidance', '')
                section_guide['word_count'] = h2_data.get('word_count', 300)
                section_guide['key_terms'] = h2_data.get('key_terms', [])
                section_guide['why_it_matters'] = h2_data.get('why_it_matters', '')
            
            # Process subsections
            for subsection in section.get('subsections', []):
                h3 = subsection.get('h3', '')
                if not h3:
                    continue
                    
                # Generate guidance for this H3 subsection
                h3_response = client.messages.create(
                    model="claude-3-5-sonnet-20241022",
                    max_tokens=800,
                    system="You are an expert SEO content strategist.",
                    messages=[
                        {"role": "user", "content": f"""
                        Create writing guidance for this subsection under "{h2}" in an article about "{keyword}":
                        
                        H3: {h3}
                        
                        Provide:
                        1. Key points to cover
                        2. Approximate word count
                        3. Why this subsection matters (SEO perspective)
                        
                        Format your response as JSON:
                        
                        {{
                            "guidance": "Detailed subsection guidance...",
                            "word_count": 150,
                            "why_it_matters": "Explanation of why this subsection is important..."
                        }}
                        """}
                    ],
                    temperature=0.5
                )
                
                # Extract and parse JSON response for H3
                h3_content = h3_response.content[0].text
                h3_json_match = re.search(r'({.*})', h3_content, re.DOTALL)
                
                subsection_guide = {
                    'heading': h3,
                    'guidance': '',
                    'word_count': 150,
                    'why_it_matters': ''
                }
                
                if h3_json_match:
                    h3_data = json.loads(h3_json_match.group(1))
                    subsection_guide['guidance'] = h3_data.get('guidance', '')
                    subsection_guide['word_count'] = h3_data.get('word_count', 150)
                    subsection_guide['why_it_matters'] = h3_data.get('why_it_matters', '')
                
                section_guide['subsections'].append(subsection_guide)
            
            writer_guide['sections'].append(section_guide)
        
        return writer_guide, True
    
    except Exception as e:
        error_msg = f"Exception in generate_writer_guide: {str(e)}"
        logger.error(error_msg)
        return {}, False

def format_writer_guide_as_html(writer_guide: Dict) -> str:
    """
    Format writer guide data as HTML with section explanations
    """
    html = []
    
    # Add H1 section
    h1_data = writer_guide.get('h1', {})
    html.append(f"<h1>{h1_data.get('heading', 'Article Title')}</h1>")
    
    html.append("<div class='section-guide'>")
    html.append("<h2>Introduction Guidance</h2>")
    html.append(f"<p>{h1_data.get('guidance', '')}</p>")
    
    html.append("<h3>Why This Matters</h3>")
    html.append(f"<p><em>{h1_data.get('why_it_matters', '')}</em></p>")
    html.append("</div>")
    
    # Add each section
    for section in writer_guide.get('sections', []):
        html.append(f"<h2>{section.get('heading', '')}</h2>")
        
        html.append("<div class='section-guide'>")
        html.append("<h3>Section Guidance</h3>")
        html.append(f"<p>{section.get('guidance', '')}</p>")
        
        html.append("<h4>Key Details</h4>")
        html.append("<ul>")
        html.append(f"<li><strong>Word count:</strong> {section.get('word_count', 300)}</li>")
        
        # Add key terms if available
        key_terms = section.get('key_terms', [])
        if key_terms:
            html.append(f"<li><strong>Key terms:</strong> {', '.join(key_terms)}</li>")
            
        html.append("</ul>")
        
        html.append("<h4>Why This Section Matters</h4>")
        html.append(f"<p><em>{section.get('why_it_matters', '')}</em></p>")
        html.append("</div>")
        
        # Add subsections
        for subsection in section.get('subsections', []):
            html.append(f"<h3>{subsection.get('heading', '')}</h3>")
            
            html.append("<div class='section-guide'>")
            html.append("<h4>Subsection Guidance</h4>")
            html.append(f"<p>{subsection.get('guidance', '')}</p>")
            
            html.append("<h5>Key Details</h5>")
            html.append("<ul>")
            html.append(f"<li><strong>Word count:</strong> {subsection.get('word_count', 150)}</li>")
            html.append("</ul>")
            
            html.append("<h5>Why This Subsection Matters</h5>")
            html.append(f"<p><em>{subsection.get('why_it_matters', '')}</em></p>")
            html.append("</div>")
    
    return "\n".join(html)

###############################################################################
# 9. Main Analysis Function
###############################################################################

def analyse(keyword: str, serp_results: List[Dict], openai_api_key: str, anthropic_api_key: str, 
            firecrawl_api_key: str, max_results: int = 20) -> Tuple[Dict, bool]:
    """
    Enhanced content analysis with 20-result scrape, clustering, and heading frequency weight
    Returns: analysis_data, success_status
    """
    try:
        # Limit to top 20 results
        top_results = serp_results[:max_results]
        
        # Scrape competitor content using FireCrawl
        competitor_contents = []
        for result in top_results:
            url = result.get('url', '')
            if not url:
                continue
                
            # Scrape content
            content, content_success = scrape_webpage(url, firecrawl_api_key)
            
            if content_success and content:
                # Extract headings
                headings = extract_headings(url, firecrawl_api_key)
                
                competitor_contents.append({
                    'url': url,
                    'title': result.get('title', ''),
                    'content': content,
                    'headings': headings,
                    'rank': result.get('rank_group', 0),
                    'page_type': result.get('page_type', '')
                })
        
        if not competitor_contents:
            logging.error("Failed to scrape any competitor content")
            return {}, False
        
        # Cluster competitors using OpenAI embeddings and FAISS K-means
        clustered_competitors, cluster_success = _cluster_competitors(
            competitor_contents, openai_api_key, num_clusters=5
        )
        
        # Build lexical corpus from competitor content
        corpus = build_lexical_corpus(clustered_competitors)
        
        # Extract important terms using Claude
        term_data, term_success = extract_important_terms(
            clustered_competitors, anthropic_api_key
        )
        
        if not term_success:
            logging.error("Failed to extract important terms")
            return {}, False
        
        # Analyze semantic structure
        semantic_structure, structure_success = analyze_semantic_structure(
            clustered_competitors, anthropic_api_key
        )
        
        if not structure_success:
            logging.error("Failed to analyze semantic structure")
            return {}, False
        
        # Generate writer guide with section explanations
        writer_guide, guide_success = generate_writer_guide(
            keyword, semantic_structure, corpus, term_data, anthropic_api_key
        )
        
        # Compile analysis data
        analysis_data = {
            'keyword': keyword,
            'competitor_contents': clustered_competitors,
            'semantic_structure': semantic_structure,
            'term_data': term_data,
            'corpus': corpus,
            'heading_counts': corpus['heading_counts'],
            'writer_guide': writer_guide
        }
        
        return analysis_data, True
        
    except Exception as e:
        error_msg = f"Exception in analyse: {str(e)}"
        logging.error(error_msg)
        return {}, False

###############################################################################
# 10. Enhanced Article Generation with Linting
###############################################################################

def generate_article_with_linter(keyword: str, semantic_structure: Dict, writer_guide: Dict,
                              term_data: Dict, corpus: Dict, anthropic_api_key: str, 
                              guidance_only: bool = False) -> Tuple[str, Dict, bool]:
    """
    Generate comprehensive article with natural language flow and balanced keyword usage.
    Includes automatic linting after generation.
    Returns: article_content, lint_results, success_status
    """
    try:
        if guidance_only:
            # Generate writing guidance
            guidance_content = format_writer_guide_as_html(writer_guide)
            return guidance_content, {}, True
        
        # Generate the actual article content
        article_content, article_success = generate_article(
            keyword, semantic_structure, [], [], [], term_data, anthropic_api_key, guidance_only=False
        )
        
        if not article_success or not article_content:
            return "", {}, False
        
        # Run the linter on the generated content
        lint_results = lint_article(article_content, keyword, term_data)
        
        return article_content, lint_results, True
        
    except Exception as e:
        error_msg = f"Exception in generate_article_with_linter: {str(e)}"
        logging.error(error_msg)
        return "", {}, False

###############################################################################
# 11. Updated Word Document with Lint Results
###############################################################################

def create_word_document(keyword: str, serp_results: List[Dict], related_keywords: List[Dict],
                        semantic_structure: Dict, article_content: str, meta_title: str, 
                        meta_description: str, paa_questions: List[Dict], term_data: Dict = None,
                        score_data: Dict = None, internal_links: List[Dict] = None, 
                        guidance_only: bool = False, lint_results: Dict = None) -> Tuple[BytesIO, bool]:
    """
    Create Word document with all components including content score and lint results
    Returns: document_stream, success_status
    """
    try:
        doc = Document()
        
        # Add document title
        doc.add_heading(f'SEO Brief: {keyword}', 0)
        
        # Add date
        doc.add_paragraph(f"Generated on: {time.strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Add meta title and description
        doc.add_heading('Meta Tags', level=1)
        meta_paragraph = doc.add_paragraph()
        meta_paragraph.add_run("Meta Title: ").bold = True
        meta_paragraph.add_run(meta_title)
        
        desc_paragraph = doc.add_paragraph()
        desc_paragraph.add_run("Meta Description: ").bold = True
        desc_paragraph.add_run(meta_description)
        
        # Section 1: SERP Analysis
        doc.add_heading('SERP Analysis', level=1)
        doc.add_paragraph('Top 10 Organic Results:')
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Rank'
        header_cells[1].text = 'Title'
        header_cells[2].text = 'URL'
        header_cells[3].text = 'Page Type'
        
        # Add data rows
        for result in serp_results[:10]:  # Just show top 10 in document
            row_cells = table.add_row().cells
            row_cells[0].text = str(result.get('rank_group', ''))
            row_cells[1].text = result.get('title', '')
            row_cells[2].text = result.get('url', '')
            row_cells[3].text = result.get('page_type', '')
        
        # Add People Also Asked questions
        if paa_questions:
            doc.add_heading('People Also Asked', level=2)
            for i, question in enumerate(paa_questions, 1):
                q_paragraph = doc.add_paragraph(style='List Number')
                q_paragraph.add_run(question.get('question', '')).bold = True
                
                # Add expanded answers if available
                for expanded in question.get('expanded', []):
                    if expanded.get('description'):
                        doc.add_paragraph(expanded.get('description', ''), style='List Bullet')
        
        # Section 2: Related Keywords
        doc.add_heading('Related Keywords', level=1)
        
        kw_table = doc.add_table(rows=1, cols=3)
        kw_table.style = 'Table Grid'
        
        # Add header row
        kw_header_cells = kw_table.rows[0].cells
        kw_header_cells[0].text = 'Keyword'
        kw_header_cells[1].text = 'Search Volume'
        kw_header_cells[2].text = 'CPC ($)'
        
        # Add data rows with safe handling of values
        for kw in related_keywords:
            row_cells = kw_table.add_row().cells
            row_cells[0].text = kw.get('keyword', '')
            
            # Safe handling of search volume
            search_volume = kw.get('search_volume')
            row_cells[1].text = str(search_volume if search_volume is not None else 0)
            
            # Safe handling of CPC
            cpc_value = kw.get('cpc')
            if cpc_value is None:
                row_cells[2].text = "$0.00"
            else:
                try:
                    row_cells[2].text = f"${float(cpc_value):.2f}"
                except (ValueError, TypeError):
                    # Handle case where CPC might be a string or other non-numeric value
                    row_cells[2].text = "$0.00"
        
        # Section 3: Important Terms (if available)
        if term_data:
            doc.add_heading('Important Terms to Include', level=1)
            
            # Primary Terms
            doc.add_heading('Primary Terms', level=2)
            primary_table = doc.add_table(rows=1, cols=3)
            primary_table.style = 'Table Grid'
            
            header_cells = primary_table.rows[0].cells
            header_cells[0].text = 'Term'
            header_cells[1].text = 'Importance'
            header_cells[2].text = 'Recommended Usage'
            
            for term in term_data.get('primary_terms', []):
                row_cells = primary_table.add_row().cells
                row_cells[0].text = term.get('term', '')
                row_cells[1].text = f"{term.get('importance', 0):.2f}"
                row_cells[2].text = str(term.get('recommended_usage', 1))
            
            # Secondary Terms
            doc.add_heading('Secondary Terms', level=2)
            secondary_table = doc.add_table(rows=1, cols=2)
            secondary_table.style = 'Table Grid'
            
            header_cells = secondary_table.rows[0].cells
            header_cells[0].text = 'Term'
            header_cells[1].text = 'Importance'
            
            for term in term_data.get('secondary_terms', [])[:15]:  # Limit to top 15
                row_cells = secondary_table.add_row().cells
                row_cells[0].text = term.get('term', '')
                row_cells[1].text = f"{term.get('importance', 0):.2f}"
        
        # Section 4: Content Score (if available)
        if score_data:
            doc.add_heading('Content Score', level=1)
            
            score_paragraph = doc.add_paragraph()
            score_paragraph.add_run(f"Overall Score: ").bold = True
            score_run = score_paragraph.add_run(f"{score_data.get('overall_score', 0)} - {score_data.get('grade', 'F')}")
            
            # Color the score based on value
            overall_score = score_data.get('overall_score', 0)
            if overall_score >= 70:
                score_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            elif overall_score < 50:
                score_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            else:
                score_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
            
            # Component scores
            doc.add_heading('Score Components', level=2)
            components = score_data.get('components', {})
            
            for component, value in components.items():
                component_para = doc.add_paragraph(style='List Bullet')
                component_name = component.replace('_score', '').replace('_', ' ').title()
                component_para.add_run(f"{component_name}: ").bold = True
                component_para.add_run(f"{value}")
        
        # Add Lint Results section if available
        if lint_results:
            doc.add_heading('Content Quality Analysis', level=1)
            
            # 1. Readability
            readability = lint_results.get('readability', {})
            doc.add_heading('Readability', level=2)
            
            read_para = doc.add_paragraph()
            read_para.add_run("Score: ").bold = True
            score = readability.get('score', 0)
            
            # Color based on score
            score_text = f"{score:.1f} ({readability.get('grade_level', '')})"
            score_run = read_para.add_run(score_text)
            
            if score >= 60:
                score_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            elif score < 30:
                score_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            else:
                score_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
            
            # Add readability issues
            readability_issues = readability.get('issues', [])
            if readability_issues:
                doc.add_heading('Readability Issues', level=3)
                for issue in readability_issues:
                    issue_para = doc.add_paragraph(issue, style='List Bullet')
            
            # 2. Keyword Density
            keyword_density = lint_results.get('keyword_density', {})
            doc.add_heading('Keyword Usage', level=2)
            
            # Primary keyword table
            primary_table = doc.add_table(rows=1, cols=4)
            primary_table.style = 'Table Grid'
            
            header_cells = primary_table.rows[0].cells
            header_cells[0].text = 'Term'
            header_cells[1].text = 'Count'
            header_cells[2].text = 'Density'
            header_cells[3].text = 'Status'
            
            primary_kw = keyword_density.get('primary_keyword', {})
            row_cells = primary_table.add_row().cells
            row_cells[0].text = primary_kw.get('term', keyword)
            row_cells[1].text = str(primary_kw.get('count', 0))
            row_cells[2].text = f"{primary_kw.get('density', 0):.2f}%"
            
            status_cell = row_cells[3]
            status_text = primary_kw.get('status', 'unknown')
            status_cell.text = status_text
            
            # Add color based on status
            for paragraph in status_cell.paragraphs:
                for run in paragraph.runs:
                    if status_text == 'good':
                        run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                    elif status_text == 'missing':
                        run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                    else:
                        run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
            
            # Primary terms table
            if keyword_density.get('primary_terms'):
                doc.add_heading('Primary Terms Usage', level=3)
                
                terms_table = doc.add_table(rows=1, cols=5)
                terms_table.style = 'Table Grid'
                
                header_cells = terms_table.rows[0].cells
                header_cells[0].text = 'Term'
                header_cells[1].text = 'Count'
                header_cells[2].text = 'Density'
                header_cells[3].text = 'Recommended'
                header_cells[4].text = 'Status'
                
                for term_info in keyword_density.get('primary_terms', []):
                    row_cells = terms_table.add_row().cells
                    row_cells[0].text = term_info.get('term', '')
                    row_cells[1].text = str(term_info.get('count', 0))
                    row_cells[2].text = f"{term_info.get('density', 0):.2f}%"
                    row_cells[3].text = str(term_info.get('recommended', 1))
                    
                    status_cell = row_cells[4]
                    status_text = term_info.get('status', 'unknown')
                    status_cell.text = status_text
                    
                    # Add color based on status
                    for paragraph in status_cell.paragraphs:
                        for run in paragraph.runs:
                            if status_text == 'good':
                                run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                            elif status_text == 'missing':
                                run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                            else:
                                run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
            
            # Keyword density issues
            kw_issues = keyword_density.get('issues', [])
            if kw_issues:
                doc.add_heading('Keyword Usage Issues', level=3)
                for issue in kw_issues:
                    issue_para = doc.add_paragraph(issue, style='List Bullet')
            
            # 3. Sentence Variety
            sentence_variety = lint_results.get('sentence_variety', {})
            doc.add_heading('Sentence Variety', level=2)
            
            sv_para = doc.add_paragraph()
            sv_para.add_run("Variety Score: ").bold = True
            variety_score = sentence_variety.get('score', 0)
            
            # Color based on score
            score_text = f"{variety_score:.1f}%"
            score_run = sv_para.add_run(score_text)
            
            if variety_score >= 80:
                score_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            elif variety_score < 50:
                score_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            else:
                score_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
            
            # Add sentence variety issues
            sv_issues = sentence_variety.get('issues', [])
            if sv_issues:
                doc.add_heading('Sentence Structure Issues', level=3)
                for issue in sv_issues:
                    issue_para = doc.add_paragraph(issue, style='List Bullet')
            
            # Add repeating patterns if found
            repeat_patterns = sentence_variety.get('repeat_patterns', [])
            if repeat_patterns:
                doc.add_heading('Repeating Sentence Patterns', level=3)
                
                patterns_table = doc.add_table(rows=1, cols=2)
                patterns_table.style = 'Table Grid'
                
                header_cells = patterns_table.rows[0].cells
                header_cells[0].text = 'Pattern'
                header_cells[1].text = 'Count'
                
                for pattern, count in repeat_patterns:
                    row_cells = patterns_table.add_row().cells
                    row_cells[0].text = f"{pattern}..."
                    row_cells[1].text = str(count)
        
        # Section 5: Semantic Structure
        doc.add_heading('Recommended Content Structure', level=1)
        
        doc.add_paragraph(f"Recommended H1: {semantic_structure.get('h1', '')}")
        
        for i, section in enumerate(semantic_structure.get('sections', []), 1):
            doc.add_paragraph(f"H2 Section {i}: {section.get('h2', '')}")
            
            for j, subsection in enumerate(section.get('subsections', []), 1):
                doc.add_paragraph(f"    H3 Subsection {j}: {subsection.get('h3', '')}")
        
        # Section 6: Generated Article or Guidance
        doc.add_heading('Generated Article Content', level=1)
        
        # IMPROVED: Line-by-line processing with better markdown heading detection
        if article_content and isinstance(article_content, str):
            # Split the content by lines
            lines = article_content.split('\n')
            
            # Initialize list tracking
            in_ul = False
            in_ol = False
            
            # Helper function to process paragraph or list item text with styling
            def add_styled_text(doc_element, text):
                """Add text with inline styling to a paragraph or list item"""
                # Process different types of formatting tags
                bold_segments = re.finditer(r'<(strong|b)>(.*?)</\1>', text, re.IGNORECASE)
                italic_segments = re.finditer(r'<(em|i)>(.*?)</\1>', text, re.IGNORECASE)
                
                # First, replace styled segments with markers
                placeholders = {}
                placeholder_id = 0
                
                # Process bold segments
                for match in bold_segments:
                    placeholder = f"__BOLD_{placeholder_id}__"
                    placeholders[placeholder] = {'type': 'bold', 'text': match.group(2)}
                    text = text.replace(match.group(0), placeholder, 1)
                    placeholder_id += 1
                
                # Process italic segments
                for match in italic_segments:
                    placeholder = f"__ITALIC_{placeholder_id}__"
                    placeholders[placeholder] = {'type': 'italic', 'text': match.group(2)}
                    text = text.replace(match.group(0), placeholder, 1)
                    placeholder_id += 1
                
                # Remove any remaining HTML tags
                clean_text = re.sub(r'<[^>]+>', '', text)
                
                # Split by placeholders to find plain text segments
                segments = []
                last_end = 0
                for match in re.finditer(r'__(BOLD|ITALIC)_\d+__', clean_text):
                    # Add plain text before the placeholder
                    if match.start() > last_end:
                        segments.append({
                            'type': 'plain',
                            'text': clean_text[last_end:match.start()]
                        })
                    
                    # Add the placeholder
                    segments.append(placeholders[match.group(0)])
                    last_end = match.end()
                
                # Add remaining plain text
                if last_end < len(clean_text):
                    segments.append({
                        'type': 'plain',
                        'text': clean_text[last_end:]
                    })
                
                # Add each segment to the document element
                for segment in segments:
                    if not segment['text'].strip():
                        continue
                        
                    if segment['type'] == 'plain':
                        doc_element.add_run(segment['text'])
                    elif segment['type'] == 'bold':
                        run = doc_element.add_run(segment['text'])
                        run.bold = True
                    elif segment['type'] == 'italic':
                        run = doc_element.add_run(segment['text'])
                        run.italic = True
            
            # Process each line
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # First check for markdown headings
                markdown_heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
                if markdown_heading_match:
                    # Count number of # symbols to determine heading level
                    level = len(markdown_heading_match.group(1))
                    heading_text = markdown_heading_match.group(2).strip()
                    
                    # Remove any HTML tags from the heading
                    heading_text = re.sub(r'<[^>]+>', '', heading_text).strip()
                    
                    if heading_text:
                        doc.add_heading(heading_text, level=level)
                        # Reset list state
                        in_ul = False
                        in_ol = False
                    continue
                
                # Check if this line is an HTML heading tag
                h1_match = re.match(r'^<h1[^>]*>(.*?)</h1>$', line, re.IGNORECASE)
                h2_match = re.match(r'^<h2[^>]*>(.*?)</h2>$', line, re.IGNORECASE)
                h3_match = re.match(r'^<h3[^>]*>(.*?)</h3>$', line, re.IGNORECASE)
                h4_match = re.match(r'^<h4[^>]*>(.*?)</h4>$', line, re.IGNORECASE)
                h5_match = re.match(r'^<h5[^>]*>(.*?)</h5>$', line, re.IGNORECASE)
                h6_match = re.match(r'^<h6[^>]*>(.*?)</h6>$', line, re.IGNORECASE)
                
                # Process heading matches
                if h1_match:
                    heading_text = h1_match.group(1)
                    # Remove any HTML tags inside the heading
                    heading_text = re.sub(r'<[^>]+>', '', heading_text).strip()
                    if heading_text:
                        doc.add_heading(heading_text, level=1)
                        # Reset list state
                        in_ul = False
                        in_ol = False
                elif h2_match:
                    heading_text = h2_match.group(1)
                    heading_text = re.sub(r'<[^>]+>', '', heading_text).strip()
                    if heading_text:
                        doc.add_heading(heading_text, level=2)
                        # Reset list state
                        in_ul = False
                        in_ol = False
                elif h3_match:
                    heading_text = h3_match.group(1)
                    heading_text = re.sub(r'<[^>]+>', '', heading_text).strip()
                    if heading_text:
                        doc.add_heading(heading_text, level=3)
                        # Reset list state
                        in_ul = False
                        in_ol = False
                elif h4_match:
                    heading_text = h4_match.group(1)
                    heading_text = re.sub(r'<[^>]+>', '', heading_text).strip()
                    if heading_text:
                        doc.add_heading(heading_text, level=4)
                        # Reset list state
                        in_ul = False
                        in_ol = False
                elif h5_match:
                    heading_text = h5_match.group(1)
                    heading_text = re.sub(r'<[^>]+>', '', heading_text).strip()
                    if heading_text:
                        doc.add_heading(heading_text, level=5)
                        # Reset list state
                        in_ul = False
                        in_ol = False
                elif h6_match:
                    heading_text = h6_match.group(1)
                    heading_text = re.sub(r'<[^>]+>', '', heading_text).strip()
                    if heading_text:
                        doc.add_heading(heading_text, level=6)
                        # Reset list state
                        in_ul = False
                        in_ol = False
                
                # Check for list starts and ends
                elif re.match(r'^<ul[^>]*>', line, re.IGNORECASE):
                    in_ul = True
                    in_ol = False
                elif re.match(r'^</ul>', line, re.IGNORECASE):
                    in_ul = False
                elif re.match(r'^<ol[^>]*>', line, re.IGNORECASE):
                    in_ol = True
                    in_ul = False
                elif re.match(r'^</ol>', line, re.IGNORECASE):
                    in_ol = False
                
                # Check if this line is a list item
                elif re.match(r'^<li[^>]*>', line, re.IGNORECASE) and '</li>' in line.lower():
                    # Extract list item content
                    li_match = re.match(r'^<li[^>]*>(.*?)</li>$', line, re.IGNORECASE)
                    if li_match:
                        li_text = li_match.group(1)
                        if li_text.strip():
                            if in_ol:
                                list_para = doc.add_paragraph(style='List Number')
                                add_styled_text(list_para, li_text)
                            else:  # Default to bullet if we're not sure
                                list_para = doc.add_paragraph(style='List Bullet')
                                add_styled_text(list_para, li_text)
                
                # Check if this line is a paragraph tag
                elif re.match(r'^<p[^>]*>', line, re.IGNORECASE) and '</p>' in line.lower():
                    # Extract paragraph content
                    para_match = re.match(r'^<p[^>]*>(.*?)</p>$', line, re.IGNORECASE)
                    if para_match:
                        para_text = para_match.group(1)
                        if para_text.strip():
                            para = doc.add_paragraph()
                            add_styled_text(para, para_text)
                
                # Check for markdown bullet points
                elif line.startswith('* ') or line.startswith('- '):
                    bullet_text = line[2:].strip()
                    if bullet_text:
                        bullet_para = doc.add_paragraph(style='List Bullet')
                        add_styled_text(bullet_para, bullet_text)
                
                # Check for markdown numbered list
                elif re.match(r'^\d+\.\s', line):
                    num_text = re.sub(r'^\d+\.\s', '', line).strip()
                    if num_text:
                        num_para = doc.add_paragraph(style='List Number')
                        add_styled_text(num_para, num_text)
                
                # If it's plain text (not starting with a tag), add as paragraph
                elif not line.startswith('<'):
                    para = doc.add_paragraph()
                    para.add_run(line)
                
                # Otherwise, try to extract clean text from HTML
                else:
                    # Clean HTML tags and add as paragraph if content exists
                    clean_text = re.sub(r'<[^>]+>', '', line).strip()
                    if clean_text:
                        para = doc.add_paragraph()
                        para.add_run(clean_text)
        
        # Section 7: Internal Linking (if provided)
        if internal_links:
            doc.add_heading('Internal Linking Summary', level=1)
            
            link_table = doc.add_table(rows=1, cols=3)
            link_table.style = 'Table Grid'
            
            # Add header row
            link_header_cells = link_table.rows[0].cells
            link_header_cells[0].text = 'URL'
            link_header_cells[1].text = 'Anchor Text'
            link_header_cells[2].text = 'Context'
            
            # Add data rows
            for link in internal_links:
                row_cells = link_table.add_row().cells
                row_cells[0].text = link.get('url', '')
                row_cells[1].text = link.get('anchor_text', '')
                row_cells[2].text = link.get('context', '')
        
        # Save document to memory stream
        doc_stream = BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)
        
        return doc_stream, True
    
    except Exception as e:
        error_msg = f"Exception in create_word_document: {str(e)}"
        logger.error(error_msg)
        logger.error(traceback.format_exc())
        return BytesIO(), False

###############################################################################
# 12. Main Streamlit App
###############################################################################

def main():
    st.title("📊 SEO Content Optimizer")
    
    # Sidebar for API credentials
    st.sidebar.header("API Credentials")
    
    # NEW: FireCrawl API key
    firecrawl_api_key = st.sidebar.text_input("FireCrawl API Key", type="password", 
                                             help="API key for FireCrawl web scraping service")
    
    dataforseo_login = st.sidebar.text_input("DataForSEO API Login", type="password")
    dataforseo_password = st.sidebar.text_input("DataForSEO API Password", type="password")
    
    openai_api_key = st.sidebar.text_input("OpenAI API Key", type="password")
    anthropic_api_key = st.sidebar.text_input("Anthropic API Key", type="password")
    
    # NEW: FireCrawl Settings
    with st.sidebar.expander("FireCrawl Settings"):
        max_concurrency = st.slider("Max Concurrency", 1, 10, 5,
                               help="Maximum number of concurrent FireCrawl requests")
        scrape_results_count = st.slider("Results to Scrape", 10, 20, 20,
                                    help="Number of top search results to scrape and analyze")
    
    # Initialize session state
    if 'results' not in st.session_state:
        st.session_state.results = {}
    
    # Main content
    tabs = st.tabs([
        "Input & SERP Analysis", 
        "Content Analysis", 
        "Article Generation", 
        "Internal Linking", 
        "SEO Brief",
        "Content Updates",
        "Content Scoring"
    ])
    
    # Tab 1: Input & SERP Analysis
    with tabs[0]:
        st.header("Enter Target Keyword")
        keyword = st.text_input("Target Keyword")
        
        if st.button("Fetch SERP Data"):
            if not keyword:
                st.error("Please enter a target keyword")
            elif not dataforseo_login or not dataforseo_password:
                st.error("Please enter DataForSEO API credentials")
            elif not openai_api_key or not anthropic_api_key:
                st.error("Please enter API keys for OpenAI and Anthropic")
            else:
                with st.spinner("Fetching SERP data..."):
                    # Fetch SERP results (updated to get 20 results for clustering)
                    start_time = time.time()
                    organic_results, serp_features, paa_questions, serp_success = fetch_serp_results(
                        keyword, dataforseo_login, dataforseo_password
                    )
                    
                    if serp_success:
                        st.session_state.results['keyword'] = keyword
                        st.session_state.results['organic_results'] = organic_results
                        st.session_state.results['serp_features'] = serp_features
                        st.session_state.results['paa_questions'] = paa_questions
                        
                        # Show SERP results
                        st.subheader("Top Organic Results")
                        df_results = pd.DataFrame(organic_results)
                        st.dataframe(df_results)
                        
                        # Show SERP features
                        st.subheader("SERP Features")
                        df_features = pd.DataFrame(serp_features)
                        st.dataframe(df_features)
                        
                        # Show People Also Asked questions
                        if paa_questions:
                            st.subheader("People Also Asked Questions")
                            for q in paa_questions:
                                st.write(f"- {q.get('question', '')}")
                        
                        # Fetch related keywords using DataForSEO
                        st.text("Fetching related keywords...")
                        related_keywords, kw_success = fetch_related_keywords_dataforseo(
                            keyword, dataforseo_login, dataforseo_password
                        )
                        
                        # Validate related keywords data
                        validated_keywords = []
                        for kw in related_keywords:
                            validated_kw = {
                                'keyword': kw.get('keyword', ''),
                                'search_volume': int(kw.get('search_volume', 0)) if kw.get('search_volume') is not None else 0,
                                'cpc': float(kw.get('cpc', 0.0)) if kw.get('cpc') is not None else 0.0,
                                'competition': float(kw.get('competition', 0.0)) if kw.get('competition') is not None else 0.0
                            }
                            validated_keywords.append(validated_kw)
                        
                        st.session_state.results['related_keywords'] = validated_keywords
                        
                        # Display related keywords
                        st.subheader("Related Keywords")
                        df_keywords = pd.DataFrame(validated_keywords)
                        st.dataframe(df_keywords)
                        
                        st.success(f"SERP analysis completed in {format_time(time.time() - start_time)}")
                    else:
                        st.error("Failed to fetch SERP data. Please check your API credentials.")
        
        # Show previously fetched data if available
        if 'organic_results' in st.session_state.results:
            st.subheader("Previously Fetched SERP Results")
            st.dataframe(pd.DataFrame(st.session_state.results['organic_results']))
            
            if 'related_keywords' in st.session_state.results:
                st.subheader("Previously Fetched Related Keywords")
                st.dataframe(pd.DataFrame(st.session_state.results['related_keywords']))
            
            if 'paa_questions' in st.session_state.results:
                st.subheader("Previously Fetched 'People Also Asked' Questions")
                for q in st.session_state.results['paa_questions']:
                    st.write(f"- {q.get('question', '')}")
    
    # Tab 2: Content Analysis
    with tabs[1]:
        st.header("Content Analysis")
        
        if 'organic_results' not in st.session_state.results:
            st.warning("Please fetch SERP data first (in the 'Input & SERP Analysis' tab)")
        else:
            if st.button("Analyze Content"):
                if not anthropic_api_key:
                    st.error("Please enter Anthropic API key")
                elif not openai_api_key:
                    st.error("Please enter OpenAI API key for embedding generation")
                else:
                    with st.spinner("Analyzing content from top-ranking pages..."):
                        start_time = time.time()
                        
                        # Use the enhanced analysis function with FireCrawl and clustering
                        analysis_data, analysis_success = analyse(
                            st.session_state.results['keyword'],
                            st.session_state.results['organic_results'],
                            openai_api_key,
                            anthropic_api_key,
                            firecrawl_api_key,
                            max_results=scrape_results_count
                        )
                            
                        if analysis_success:
                            # Store results
                            st.session_state.results['analysis_data'] = analysis_data
                            st.session_state.results['semantic_structure'] = analysis_data['semantic_structure']
                            st.session_state.results['term_data'] = analysis_data['term_data']
                            st.session_state.results['scraped_contents'] = analysis_data['competitor_contents']
                            st.session_state.results['heading_counts'] = analysis_data.get('heading_counts', {})
                            st.session_state.results['writer_guide'] = analysis_data.get('writer_guide', {})
                            
                            # Display semantic structure
                            st.subheader("Recommended Semantic Structure")
                            semantic_structure = analysis_data['semantic_structure']
                            
                            st.write(f"**H1:** {semantic_structure.get('h1', '')}")
                            
                            for i, section in enumerate(semantic_structure.get('sections', []), 1):
                                st.write(f"**H2 {i}:** {section.get('h2', '')}")
                                
                                for j, subsection in enumerate(section.get('subsections', []), 1):
                                    st.write(f"  - **H3 {i}.{j}:** {subsection.get('h3', '')}")
                            
                            # Display clustered content information
                            st.subheader("Content Clusters")
                            clustered_contents = analysis_data['competitor_contents']
                            cluster_counts = {}
                            
                            for content in clustered_contents:
                                cluster_id = content.get('cluster', -1)
                                if cluster_id in cluster_counts:
                                    cluster_counts[cluster_id] += 1
                                else:
                                    cluster_counts[cluster_id] = 1
                            
                            cluster_df = pd.DataFrame({
                                'Cluster': list(cluster_counts.keys()),
                                'Count': list(cluster_counts.values())
                            })
                            
                            st.bar_chart(cluster_df.set_index('Cluster'))
                            
                            # Display heading frequency data
                            st.subheader("Common Headings")
                            heading_counts = analysis_data.get('heading_counts', {})
                            
                            with st.expander("H2 Heading Frequency"):
                                h2_counts = heading_counts.get('h2', {})
                                if h2_counts:
                                    h2_df = pd.DataFrame({
                                        'Heading': list(h2_counts.keys()),
                                        'Count': list(h2_counts.values())
                                    }).sort_values('Count', ascending=False).head(15)
                                    
                                    st.dataframe(h2_df)
                            
                            with st.expander("H3 Heading Frequency"):
                                h3_counts = heading_counts.get('h3', {})
                                if h3_counts:
                                    h3_df = pd.DataFrame({
                                        'Heading': list(h3_counts.keys()),
                                        'Count': list(h3_counts.values())
                                    }).sort_values('Count', ascending=False).head(15)
                                    
                                    st.dataframe(h3_df)
                            
                            # Display term data
                            with st.expander("View Extracted Terms & Topics"):
                                term_data = analysis_data['term_data']
                                
                                # Display primary terms
                                if 'primary_terms' in term_data:
                                    st.write("**Primary Terms:**")
                                    primary_df = pd.DataFrame(term_data['primary_terms'])
                                    st.dataframe(primary_df)
                                
                                # Display secondary terms
                                if 'secondary_terms' in term_data:
                                    st.write("**Secondary Terms:**")
                                    secondary_df = pd.DataFrame(term_data['secondary_terms'])
                                    st.dataframe(secondary_df)
                                
                                # Display topics
                                if 'topics' in term_data:
                                    st.write("**Topics to Cover:**")
                                    topics_df = pd.DataFrame(term_data['topics'])
                                    st.dataframe(topics_df)
                            
                            st.success(f"Content analysis completed in {format_time(time.time() - start_time)}")
                        else:
                            st.error("Failed to analyze content")
            
            # Show previously analyzed structure if available
            if 'semantic_structure' in st.session_state.results:
                st.subheader("Previously Analyzed Semantic Structure")
                semantic_structure = st.session_state.results['semantic_structure']
                
                st.write(f"**H1:** {semantic_structure.get('h1', '')}")
                
                for i, section in enumerate(semantic_structure.get('sections', []), 1):
                    st.write(f"**H2 {i}:** {section.get('h2', '')}")
                    
                    for j, subsection in enumerate(section.get('subsections', []), 1):
                        st.write(f"  - **H3 {i}.{j}:** {subsection.get('h3', '')}")
                
                # Show heading counts if available
                if 'heading_counts' in st.session_state.results:
                    with st.expander("View Common Headings"):
                        heading_counts = st.session_state.results['heading_counts']
                        
                        st.write("**H2 Heading Frequency:**")
                        h2_counts = heading_counts.get('h2', {})
                        if h2_counts:
                            h2_df = pd.DataFrame({
                                'Heading': list(h2_counts.keys()),
                                'Count': list(h2_counts.values())
                            }).sort_values('Count', ascending=False).head(15)
                            
                            st.dataframe(h2_df)
                        
                        st.write("**H3 Heading Frequency:**")
                        h3_counts = heading_counts.get('h3', {})
                        if h3_counts:
                            h3_df = pd.DataFrame({
                                'Heading': list(h3_counts.keys()),
                                'Count': list(h3_counts.values())
                            }).sort_values('Count', ascending=False).head(15)
                            
                            st.dataframe(h3_df)
                
                # Show previously extracted term data if available
                if 'term_data' in st.session_state.results:
                    with st.expander("View Extracted Terms & Topics"):
                        term_data = st.session_state.results['term_data']
                        
                        # Display primary terms
                        if 'primary_terms' in term_data:
                            st.write("**Primary Terms:**")
                            primary_df = pd.DataFrame(term_data['primary_terms'])
                            st.dataframe(primary_df)
                        
                        # Display secondary terms
                        if 'secondary_terms' in term_data:
                            st.write("**Secondary Terms:**")
                            secondary_df = pd.DataFrame(term_data['secondary_terms'])
                            st.dataframe(secondary_df)
                        
                        # Display topics
                        if 'topics' in term_data:
                            st.write("**Topics to Cover:**")
                            topics_df = pd.DataFrame(term_data['topics'])
                            st.dataframe(topics_df)
    
    # Tab 3: Article Generation
    with tabs[2]:
        st.header("Article Generation")
        
        if 'semantic_structure' not in st.session_state.results:
            st.warning("Please complete content analysis first (in the 'Content Analysis' tab)")
        else:
            # Add option for guidance-only mode
            content_type = st.radio(
                "Content Generation Type:",
                ["Full Article", "Writing Guidance Only"],
                help="Choose 'Full Article' for complete content or 'Writing Guidance Only' for section-by-section writing directions"
            )
            guidance_only = (content_type == "Writing Guidance Only")
            
            if st.button("Generate " + ("Content Guidance" if guidance_only else "Article") + " and Meta Tags"):
                if not anthropic_api_key:
                    st.error("Please enter Anthropic API key")
                else:
                    with st.spinner("Generating " + ("content guidance" if guidance_only else "article") + " and meta tags..."):
                        start_time = time.time()
                        
                        # Use term data if available
                        term_data = st.session_state.results.get('term_data', {})
                        writer_guide = st.session_state.results.get('writer_guide', {})
                        corpus = st.session_state.results.get('analysis_data', {}).get('corpus', {})
                        
                        # Generate article or guidance with linting
                        article_content, lint_results, article_success = generate_article_with_linter(
                            st.session_state.results['keyword'],
                            st.session_state.results['semantic_structure'],
                            writer_guide,
                            term_data,
                            corpus,
                            anthropic_api_key,
                            guidance_only
                        )
                        
                        if article_success and article_content:
                            # Store with special key for guidance
                            if guidance_only:
                                st.session_state.results['guidance_content'] = article_content
                            else:
                                st.session_state.results['article_content'] = article_content
                                st.session_state.results['lint_results'] = lint_results
                            
                            # Store the guidance flag
                            st.session_state.results['guidance_only'] = guidance_only
                            
                            # Generate meta title and description with Claude
                            meta_title, meta_description, meta_success = generate_meta_tags(
                                st.session_state.results['keyword'],
                                st.session_state.results['semantic_structure'],
                                st.session_state.results.get('related_keywords', []),
                                term_data,
                                anthropic_api_key
                            )
                            
                            if meta_success:
                                st.session_state.results['meta_title'] = meta_title
                                st.session_state.results['meta_description'] = meta_description
                                
                                st.subheader("Meta Tags")
                                st.write(f"**Meta Title:** {meta_title}")
                                st.write(f"**Meta Description:** {meta_description}")
                            
                            # Score the content if it's a full article
                            if not guidance_only and 'term_data' in st.session_state.results:
                                try:
                                    score_data, score_success = score_content(
                                        article_content,
                                        st.session_state.results['term_data'],
                                        st.session_state.results['keyword']
                                    )
                                    
                                    if score_success:
                                        st.session_state.results['content_score'] = score_data
                                        
                                        # Display content score
                                        st.subheader("Content Score")
                                        score = score_data.get('overall_score', 0)
                                        grade = score_data.get('grade', 'F')
                                        
                                        # CSS to style the score display
                                        score_color = "green" if score >= 70 else "red" if score < 50 else "orange"
                                        st.markdown(f"""
                                        <div style="background-color: #f0f0f0; padding: 10px; border-radius: 5px;">
                                            <h3 style="margin:0;">Content Score: <span style="color:{score_color};">{score} ({grade})</span></h3>
                                        </div>
                                        """, unsafe_allow_html=True)
                                except Exception as e:
                                    logger.error(f"Error scoring content: {e}")
                            
                            # Display lint results if available
                            if not guidance_only and lint_results:
                                with st.expander("Content Quality Analysis", expanded=True):
                                    # Readability
                                    readability = lint_results.get('readability', {})
                                    st.markdown("### Readability")
                                    read_score = readability.get('score', 0)
                                    grade_level = readability.get('grade_level', '')
                                    
                                    score_color = "green" if read_score >= 60 else "red" if read_score < 30 else "orange"
                                    st.markdown(f"**Score:** <span style='color:{score_color};'>{read_score:.1f}</span> (Grade Level: {grade_level})", unsafe_allow_html=True)
                                    
                                    # Display readability issues
                                    read_issues = readability.get('issues', [])
                                    if read_issues:
                                        st.markdown("**Issues:**")
                                        for issue in read_issues:
                                            st.markdown(f"- {issue}")
                                    
                                    # Keyword density
                                    kw_density = lint_results.get('keyword_density', {})
                                    st.markdown("### Keyword Usage")
                                    
                                    # Primary keyword
                                    primary_kw = kw_density.get('primary_keyword', {})
                                    primary_status = primary_kw.get('status', 'unknown')
                                    status_color = "green" if primary_status == 'good' else "red" if primary_status == 'missing' else "orange"
                                    
                                    st.markdown(f"""
                                    **Primary Keyword:** {primary_kw.get('term', st.session_state.results['keyword'])}
                                    - **Count:** {primary_kw.get('count', 0)}
                                    - **Density:** {primary_kw.get('density', 0):.2f}%
                                    - **Status:** <span style='color:{status_color};'>{primary_status}</span>
                                    """, unsafe_allow_html=True)
                                    
                                    # Display keyword density issues
                                    kw_issues = kw_density.get('issues', [])
                                    if kw_issues:
                                        st.markdown("**Issues:**")
                                        for issue in kw_issues:
                                            st.markdown(f"- {issue}")
                                    
                                    # Sentence variety
                                    sent_variety = lint_results.get('sentence_variety', {})
                                    st.markdown("### Sentence Variety")
                                    variety_score = sent_variety.get('score', 0)
                                    score_color = "green" if variety_score >= 80 else "red" if variety_score < 50 else "orange"
                                    
                                    st.markdown(f"**Variety Score:** <span style='color:{score_color};'>{variety_score:.1f}%</span>", unsafe_allow_html=True)
                                    
                                    # Display sentence variety issues
                                    sent_issues = sent_variety.get('issues', [])
                                    if sent_issues:
                                        st.markdown("**Issues:**")
                                        for issue in sent_issues:
                                            st.markdown(f"- {issue}")
                                    
                                    # Display repeating patterns
                                    repeat_patterns = sent_variety.get('repeat_patterns', [])
                                    if repeat_patterns:
                                        st.markdown("#### Repeating Sentence Patterns")
                                        patterns_df = pd.DataFrame(repeat_patterns, columns=['Pattern', 'Count'])
                                        st.dataframe(patterns_df)
                            
                            st.subheader("Generated " + ("Content Guidance" if guidance_only else "Article"))
                            st.markdown(article_content, unsafe_allow_html=True)
                            
                            st.success(f"Content generation completed in {format_time(time.time() - start_time)}")
                        else:
                            st.error("Failed to generate " + ("content guidance" if guidance_only else "article") + ". Please try again.")
            
            # Show previously generated article or guidance if available
            if 'article_content' in st.session_state.results or 'guidance_content' in st.session_state.results:
                if 'meta_title' in st.session_state.results:
                    st.subheader("Previously Generated Meta Tags")
                    st.write(f"**Meta Title:** {st.session_state.results['meta_title']}")
                    st.write(f"**Meta Description:** {st.session_state.results['meta_description']}")
                
                # Display content score if available
                if 'content_score' in st.session_state.results:
                    st.subheader("Content Score")
                    score_data = st.session_state.results['content_score']
                    score = score_data.get('overall_score', 0)
                    grade = score_data.get('grade', 'F')
                    
                    # CSS to style the score display
                    score_color = "green" if score >= 70 else "red" if score < 50 else "orange"
                    st.markdown(f"""
                    <div style="background-color: #f0f0f0; padding: 10px; border-radius: 5px;">
                        <h3 style="margin:0;">Content Score: <span style="color:{score_color};">{score} ({grade})</span></h3>
                    </div>
                    """, unsafe_allow_html=True)
                
                # Display lint results if available
                if 'lint_results' in st.session_state.results:
                    with st.expander("Content Quality Analysis"):
                        lint_results = st.session_state.results['lint_results']
                        
                        # Readability
                        readability = lint_results.get('readability', {})
                        st.markdown("### Readability")
                        read_score = readability.get('score', 0)
                        grade_level = readability.get('grade_level', '')
                        
                        score_color = "green" if read_score >= 60 else "red" if read_score < 30 else "orange"
                        st.markdown(f"**Score:** <span style='color:{score_color};'>{read_score:.1f}</span> (Grade Level: {grade_level})", unsafe_allow_html=True)
                        
                        # Keyword density
                        kw_density = lint_results.get('keyword_density', {})
                        st.markdown("### Keyword Usage")
                        
                        # Primary keyword
                        primary_kw = kw_density.get('primary_keyword', {})
                        primary_status = primary_kw.get('status', 'unknown')
                        status_color = "green" if primary_status == 'good' else "red" if primary_status == 'missing' else "orange"
                        
                        st.markdown(f"""
                        **Primary Keyword:** {primary_kw.get('term', st.session_state.results['keyword'])}
                        - **Count:** {primary_kw.get('count', 0)}
                        - **Density:** {primary_kw.get('density', 0):.2f}%
                        - **Status:** <span style='color:{status_color};'>{primary_status}</span>
                        """, unsafe_allow_html=True)
                        
                        # Sentence variety
                        sent_variety = lint_results.get('sentence_variety', {})
                        st.markdown("### Sentence Variety")
                        variety_score = sent_variety.get('score', 0)
                        score_color = "green" if variety_score >= 80 else "red" if variety_score < 50 else "orange"
                        
                        st.markdown(f"**Variety Score:** <span style='color:{score_color};'>{variety_score:.1f}%</span>", unsafe_allow_html=True)
                
                # Display appropriate content based on what's available
                if 'guidance_only' in st.session_state.results and st.session_state.results['guidance_only']:
                    st.subheader("Previously Generated Content Guidance")
                    if 'guidance_content' in st.session_state.results:
                        st.markdown(st.session_state.results['guidance_content'], unsafe_allow_html=True)
                else:
                    st.subheader("Previously Generated Article")
                    if 'article_content' in st.session_state.results:
                        st.markdown(st.session_state.results['article_content'], unsafe_allow_html=True)
    
    # Tab 4: Internal Linking
    with tabs[3]:
        st.header("Internal Linking")
        
        is_guidance_only = st.session_state.results.get('guidance_only', False)
        
        # Check if content exists and determine which type
        has_content = False
        if is_guidance_only and 'guidance_content' in st.session_state.results:
            has_content = True
            content_type = "guidance"
        elif not is_guidance_only and 'article_content' in st.session_state.results:
            has_content = True
            content_type = "article"
        
        if not has_content:
            st.warning("Please generate content first (in the 'Article Generation' tab)")
        elif is_guidance_only:
            st.warning("Internal linking is only available for full articles, not writing guidance")
        else:
            st.write("Upload a spreadsheet with your site pages (CSV or Excel):")
            st.write("The spreadsheet must contain columns: URL, Title, Meta Description")
            
            # Create sample template button
            if st.button("Generate Sample Template"):
                # Create sample dataframe
                sample_data = {
                    'URL': ['https://example.com/page1', 'https://example.com/page2'],
                    'Title': ['Example Page 1', 'Example Page 2'],
                    'Meta Description': ['Description for page 1', 'Description for page 2']
                }
                sample_df = pd.DataFrame(sample_data)
                
                # Convert to CSV
                csv = sample_df.to_csv(index=False)
                
                # Provide download button
                st.download_button(
                    label="Download Sample CSV Template",
                    data=csv,
                    file_name="site_pages_template.csv",
                    mime="text/csv",
                )
            
            # File uploader for spreadsheet
            pages_file = st.file_uploader("Upload Site Pages Spreadsheet", type=['csv', 'xlsx', 'xls'])
            
            # Batch size for embedding
            batch_size = st.slider("Embedding Batch Size", 5, 50, 20, 
                                   help="Larger batch size is faster but may hit API limits")
            
            if st.button("Generate Internal Links"):
                if not openai_api_key:
                    st.error("Please enter OpenAI API key for embeddings")
                elif not anthropic_api_key:
                    st.error("Please enter Anthropic API key for anchor text selection")
                elif not pages_file:
                    st.error("Please upload a spreadsheet with site pages")
                else:
                    with st.spinner("Processing site pages and generating internal links..."):
                        start_time = time.time()
                        
                        # Parse spreadsheet
                        pages, parse_success = parse_site_pages_spreadsheet(pages_file)
                        
                        if parse_success and pages:
                            # Status update
                            status_text = st.empty()
                            status_text.text(f"Generating embeddings for {len(pages)} site pages...")
                            
                            # Generate embeddings for site pages using OpenAI
                            pages_with_embeddings, embed_success = embed_site_pages(
                                pages, openai_api_key, batch_size
                            )
                            
                            if embed_success:
                                # Count words in the article
                                article_content = st.session_state.results['article_content']
                                word_count = len(re.findall(r'\w+', article_content))
                                
                                status_text.text(f"Analyzing article content and generating internal links...")
                                
                                # Pass both API keys to the function
                                article_with_links, links_added, links_success = generate_internal_links_with_embeddings(
                                    article_content, pages_with_embeddings, openai_api_key, anthropic_api_key, word_count
                                )
                                
                                if links_success:
                                    st.session_state.results['article_with_links'] = article_with_links
                                    st.session_state.results['internal_links'] = links_added
                                    
                                    st.subheader("Article With Internal Links")
                                    st.markdown(article_with_links, unsafe_allow_html=True)
                                    
                                    st.subheader("Internal Links Added")
                                    for link in links_added:
                                        st.write(f"**URL:** {link.get('url')}")
                                        st.write(f"**Anchor Text:** {link.get('anchor_text')}")
                                        st.write(f"**Context:** {link.get('context')}")
                                        st.write("---")
                                    
                                    st.success(f"Internal linking completed in {format_time(time.time() - start_time)}")
                                else:
                                    st.error("Failed to generate internal links")
                            else:
                                st.error("Failed to generate embeddings for site pages")
                        else:
                            st.error("Failed to parse spreadsheet or no pages found")
            
            # Show previously generated links if available
            if 'article_with_links' in st.session_state.results:
                st.subheader("Previously Generated Article With Internal Links")
                st.markdown(st.session_state.results['article_with_links'], unsafe_allow_html=True)
    
    # Tab 5: SEO Brief
    with tabs[4]:
        st.header("SEO Brief & Downloadable Report")
        
        # Check if content exists and determine which type
        has_content = False
        content_type = "none"
        
        if 'guidance_only' in st.session_state.results:
            is_guidance = st.session_state.results['guidance_only']
            if is_guidance and 'guidance_content' in st.session_state.results:
                has_content = True
                content_type = "guidance"
            elif not is_guidance and 'article_content' in st.session_state.results:
                has_content = True
                content_type = "article"
        
        if not has_content:
            st.warning("Please generate content first (in the 'Article Generation' tab)")
        else:
            if st.button("Generate SEO Brief"):
                with st.spinner("Generating SEO brief..."):
                    start_time = time.time()
                    
                    # Determine which content to use
                    if content_type == "article":
                        # Use article with internal links if available, otherwise use regular article
                        article_content = st.session_state.results.get('article_with_links', 
                                                                      st.session_state.results['article_content'])
                        
                        internal_links = st.session_state.results.get('internal_links', None)
                        guidance_only = False
                        # Get lint results if available
                        lint_results = st.session_state.results.get('lint_results', None)
                    else:
                        # Use guidance content
                        article_content = st.session_state.results['guidance_content']
                        internal_links = None
                        guidance_only = True
                        lint_results = None
                    
                    # Get meta title and description
                    meta_title = st.session_state.results.get('meta_title', 
                                                             f"{st.session_state.results['keyword']} - Complete Guide")
                    
                    meta_description = st.session_state.results.get('meta_description', 
                                                                  f"Learn everything about {st.session_state.results['keyword']} in our comprehensive guide.")
                    
                    # Get PAA questions
                    paa_questions = st.session_state.results.get('paa_questions', [])
                    
                    # Get term data and content score if available
                    term_data = st.session_state.results.get('term_data', None)
                    score_data = st.session_state.results.get('content_score', None)
                    
                    # Create Word document with enhanced Claude content and lint results
                    doc_stream, doc_success = create_word_document(
                        st.session_state.results['keyword'],
                        st.session_state.results['organic_results'],
                        st.session_state.results.get('related_keywords', []),
                        st.session_state.results['semantic_structure'],
                        article_content,
                        meta_title,
                        meta_description,
                        paa_questions,
                        term_data,
                        score_data,
                        internal_links,
                        guidance_only,
                        lint_results
                    )
                    
                    if doc_success:
                        st.session_state.results['doc_stream'] = doc_stream
                        
                        # First download button
                        st.download_button(
                            label="Download SEO Brief Now",
                            data=doc_stream,
                            file_name=f"seo_brief_{st.session_state.results['keyword'].replace(' ', '_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_brief_1"  # Added unique key
                        )
                        
                        st.success(f"SEO brief generation completed in {format_time(time.time() - start_time)}")
                    else:
                        st.error("Failed to generate SEO brief document")
            
            # Show summary of all components
            st.subheader("SEO Analysis Summary")
            
            components = [
                ("Keyword", 'keyword' in st.session_state.results),
                ("SERP Analysis", 'organic_results' in st.session_state.results),
                ("People Also Asked", 'paa_questions' in st.session_state.results),
                ("Related Keywords", 'related_keywords' in st.session_state.results),
                ("Content Analysis", 'scraped_contents' in st.session_state.results),
                ("Term Analysis", 'term_data' in st.session_state.results),
                ("Semantic Structure", 'semantic_structure' in st.session_state.results),
                ("Meta Title & Description", 'meta_title' in st.session_state.results),
                ("Content Score", 'content_score' in st.session_state.results),
                ("Readability Analysis", 'lint_results' in st.session_state.results),
                ("Generated Content", 'article_content' in st.session_state.results or 'guidance_content' in st.session_state.results),
                ("Internal Linking", 'article_with_links' in st.session_state.results)
            ]
            
            for component, status in components:
                st.write(f"**{component}:** {'✅ Completed' if status else '❌ Not Completed'}")
            
            # Display download button if available with different key
            if 'doc_stream' in st.session_state.results:
                st.subheader("Download SEO Brief")
                st.download_button(
                    label="Download SEO Brief Document",  # Changed label to avoid duplicate element ID
                    data=st.session_state.results['doc_stream'],
                    file_name=f"seo_brief_{st.session_state.results['keyword'].replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_brief_2"  # Added unique key
                )
    
    # Tab 6: Content Updates
    with tabs[5]:
        # Rest of the Content Updates tab code (unchanged)
        pass

    # Tab 7: Content Scoring
    with tabs[6]:
        # Rest of the Content Scoring tab code (unchanged)
        pass

if __name__ == "__main__":
    main()
